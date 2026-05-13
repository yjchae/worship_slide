"""예배 슬라이드 사용자 설명서 Word 파일 생성 스크립트."""

import io
import math
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import copy

# ── 폰트 (Pretendard — 한글 지원) ────────────────────────────────────────────
_FONT_DIR = '/home/user/make_ppt/assets/fonts'
_BOLD_PATH = f'{_FONT_DIR}/Pretendard-Bold.ttf'
_REG_PATH  = f'{_FONT_DIR}/Pretendard-Regular.ttf'

def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = _BOLD_PATH if bold else _REG_PATH
    return ImageFont.truetype(path, size)

# 자주 쓰는 크기 미리 로드
_F10  = _font(10)
_F11  = _font(11)
_F12  = _font(12)
_F13  = _font(13)
_F14  = _font(14)
_F16  = _font(16)
_F18  = _font(18)
_F10B = _font(10, bold=True)
_F12B = _font(12, bold=True)
_F14B = _font(14, bold=True)
_F16B = _font(16, bold=True)
_F18B = _font(18, bold=True)


# ── 색상 팔레트 ──────────────────────────────────────────────────────────────
DARK_BG      = RGBColor(0x14, 0x36, 0x42)   # 헤더 다크 그린
TEAL         = RGBColor(0x0F, 0x8B, 0x8D)   # 헤더 밝은 틸
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
MEDIUM_GRAY  = RGBColor(0xE0, 0xE0, 0xE0)
TEXT_DARK    = RGBColor(0x1B, 0x1B, 0x1B)
ACCENT_BLUE  = RGBColor(0x0F, 0x4C, 0x5C)
ACCENT_TEAL  = RGBColor(0x0F, 0x8B, 0x8D)
NOTE_BG      = RGBColor(0xE8, 0xF4, 0xF8)
WARN_BG      = RGBColor(0xFF, 0xF3, 0xE0)
TIP_BG       = RGBColor(0xE8, 0xF5, 0xE9)


# ── 유틸 함수 ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """표 셀 배경색 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """셀 테두리 설정."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), kwargs[side].get('val', 'single'))
            border.set(qn('w:sz'), str(kwargs[side].get('sz', 4)))
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), kwargs[side].get('color', '000000'))
            tcBorders.append(border)
    tcPr.append(tcBorders)


def add_colored_run(para, text, bold=False, color: RGBColor = None,
                    size_pt=11, italic=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color
    return run


def set_para_spacing(para, before=0, after=0, line_spacing=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line_spacing:
        spacing.set(qn('w:line'), str(line_spacing))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def set_para_shading(para, hex_color: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


# ── 모의 UI 이미지 생성 ──────────────────────────────────────────────────────

def make_main_ui_image() -> Image.Image:
    """앱 메인 화면 모의 UI 이미지."""
    W, H = 1000, 620
    img = Image.new('RGB', (W, H), (245, 245, 248))
    d = ImageDraw.Draw(img)

    fB = _font(15, bold=True)
    fR = _font(13)
    fS = _font(11)

    # ── 상단 헤더 바 (그라디언트) ──
    for i in range(60):
        ratio = i / 60
        r = int(20 + (15 - 20) * ratio)
        g = int(54 + (139 - 54) * ratio)
        b = int(66 + (141 - 66) * ratio)
        d.line([(0, i), (W, i)], fill=(r, g, b))

    d.text((16, 16), "예배 슬라이드 보관함", fill=(255, 255, 255), font=fB)
    d.rounded_rectangle([260, 14, 330, 42], radius=4,
                         fill=(255, 255, 255, 46))
    d.text((270, 20), "저장 42", fill=(255, 255, 255), font=fR)

    d.rounded_rectangle([700, 12, 845, 44], radius=5,
                         fill=(255, 255, 255, 56))
    d.text((714, 20), "찬양폴더 선택", fill=(255, 255, 255), font=fR)

    d.rounded_rectangle([855, 12, 990, 44], radius=5,
                         fill=(255, 255, 255, 56))
    d.text((866, 20), "성경 불러오기", fill=(255, 255, 255), font=fR)

    # ── 발표 컨트롤 바 ──
    d.rounded_rectangle([4, 66, W - 4, 104], radius=6, fill=(228, 232, 238))
    d.text((18, 78), "발표 화면", fill=(70, 75, 90), font=fR)
    d.rounded_rectangle([W - 140, 72, W - 10, 98], radius=5,
                         fill=(65, 120, 175))
    d.text((W - 128, 78), "▶  발표 시작", fill=(255, 255, 255), font=fR)

    # ── 왼쪽 패널 (선택한 순서) ──
    d.rounded_rectangle([4, 110, 610, 280], radius=8, fill=(255, 255, 255),
                         outline=(215, 215, 218))
    d.text((16, 120), "선택한 순서  3개", fill=(25, 25, 25), font=fB)
    d.line([(16, 142), (598, 142)], fill=(225, 225, 228))

    items = [
        ("1", "주님 한 분만으로", "주님 한 분만으로 내게는 족해", False),
        ("2", "주의 이름 높이세", "온 땅이여 주께 소리 질러", False),
        ("3", "로마서 8:28", "우리가 알거니와 하나님을 사랑하는", True),
    ]
    for i, (num, title, sub, is_bible) in enumerate(items):
        y = 148 + i * 43
        if i == 0:
            d.rectangle([5, y, 608, y + 40], fill=(237, 246, 255))
        d.text((16, y + 10), num, fill=(50, 50, 50), font=fR)
        if is_bible:
            d.rounded_rectangle([40, y + 8, 78, y + 28], radius=3,
                                 fill=(187, 222, 251))
            d.text((44, y + 11), "성경", fill=(21, 101, 192), font=fS)
            d.text((84, y + 10), title, fill=(35, 35, 35), font=fR)
        else:
            d.text((40, y + 10), title, fill=(35, 35, 35), font=fR)
        d.text((40, y + 26), sub, fill=(150, 150, 155), font=fS)
        d.text((570, y + 10), "✕", fill=(180, 180, 185), font=fR)
        d.line([(6, y + 41), (607, y + 41)], fill=(235, 235, 238))

    # ── 왼쪽 패널 (검색) ──
    d.rounded_rectangle([4, 286, 610, H - 4], radius=8, fill=(255, 255, 255),
                         outline=(215, 215, 218))
    d.text((16, 296), "찬양 검색", fill=(20, 20, 20), font=fB)
    d.line([(16, 316), (105, 316)], fill=(15, 139, 141), width=2)
    d.text((126, 296), "성경 본문", fill=(140, 140, 145), font=fR)

    d.rounded_rectangle([10, 326, 598, 356], radius=5,
                         outline=(185, 185, 190), fill=(252, 252, 252))
    d.text((20, 334), "곡명 또는 가사로 검색…", fill=(185, 185, 190), font=fR)

    songs = [
        ("나는 예배자입니다", False),
        ("이 땅의 모든 것", False),
        ("주님 한 분만으로", True),
        ("주의 이름 높이세", True),
        ("찬양하라 내 영혼아", False),
    ]
    for i, (s, is_sel) in enumerate(songs):
        y = 362 + i * 44
        bg = (232, 247, 247) if is_sel else (255, 255, 255)
        d.rectangle([5, y, 608, y + 40], fill=bg)
        prefix = "✓  " if is_sel else "    "
        tc = (18, 88, 100) if is_sel else (50, 50, 55)
        d.text((18, y + 12), prefix + s, fill=tc, font=fR)
        d.line([(6, y + 40), (607, y + 40)], fill=(238, 238, 240))

    # ── 오른쪽 패널 (디자인) ──
    d.rounded_rectangle([620, 110, W - 4, H - 4], radius=8,
                         fill=(255, 255, 255), outline=(215, 215, 218))
    d.text((634, 120), "PPTX 디자인", fill=(25, 25, 25), font=fB)
    d.line([(634, 142), (988, 142)], fill=(220, 220, 222))

    # 미리보기 슬라이드
    d.rounded_rectangle([630, 150, 990, 320], radius=5, fill=(27, 27, 27))
    d.text((690, 215), "주님 한 분만으로", fill=(255, 255, 255),
           font=_font(18, bold=True))
    d.text((690, 244), "Only You, Lord", fill=(255, 241, 118), font=_font(14))
    d.text((634, 324), "▲ 첫 번째 곡 · 첫 페이지 실시간 미리보기",
           fill=(130, 130, 135), font=fS)

    # 스타일 옵션
    options = [
        ("배경색",    "■  #1B1B1B"),
        ("글자색",    "■  #FFFFFF"),
        ("글자 크기", "30 pt"),
        ("글자 위치", "중단  ▾"),
        ("영어 포함", "✓  켜짐"),
        ("곡 제목",   "꺼짐"),
    ]
    for i, (label, val) in enumerate(options):
        y = 346 + i * 30
        d.text((634, y), label, fill=(105, 105, 110), font=fR)
        d.text((760, y), val, fill=(35, 35, 40), font=fR)

    d.rounded_rectangle([630, 540, 990, 574], radius=6, fill=(15, 139, 141))
    d.text((768, 549), "PPTX 저장", fill=(255, 255, 255), font=fB)

    return img


def make_import_image() -> Image.Image:
    """찬양 가져오기 흐름 다이어그램."""
    W, H = 900, 200
    img = Image.new('RGB', (W, H), (242, 246, 250))
    d = ImageDraw.Draw(img)

    fB = _font(14, bold=True)
    fR = _font(12)

    steps = [
        ("① 찬양폴더 선택", "버튼 클릭"),
        ("② 폴더 재귀 탐색", ".pptx / .ppt 발견"),
        ("③ 가사 추출", "한/영 자동 분리"),
        ("④ DB 저장 완료", "검색 가능"),
    ]
    colors = [(20, 54, 66), (15, 95, 110), (15, 139, 141), (14, 155, 100)]
    box_w = 190

    for i, ((line1, line2), color) in enumerate(zip(steps, colors)):
        x0 = 20 + i * (box_w + 24)
        x1 = x0 + box_w
        d.rounded_rectangle([x0, 20, x1, 155], radius=10, fill=color)
        d.text((x0 + 12, 55), line1, fill=(255, 255, 255), font=fB)
        d.text((x0 + 12, 80), line2, fill=(195, 230, 230), font=fR)
        if i < 3:
            mx = x1 + 12
            d.polygon([(mx, 82), (mx + 14, 92), (mx, 102)],
                       fill=(100, 148, 155))

    d.text((310, 170), "중복 곡은 자동으로 건너뜁니다",
           fill=(100, 118, 128), font=fR)

    return img


def make_staging_image() -> Image.Image:
    """선택한 순서(스테이징) 패널 이미지."""
    W, H = 560, 310
    img = Image.new('RGB', (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)

    fB = _font(14, bold=True)
    fR = _font(13)
    fS = _font(11)

    d.rectangle([0, 0, W, 44], fill=(244, 244, 247))
    d.text((16, 12), "선택한 순서  3개", fill=(25, 25, 30), font=fB)
    d.line([(0, 44), (W, 44)], fill=(218, 218, 222))

    rows = [
        ("1", "주님 한 분만으로", "주님 한 분만으로 내게는 족해", False),
        ("2", "주의 이름 높이세", "온 땅이여 주께 소리 질러", False),
        ("3", "로마서 8:28", "우리가 알거니와", True),
    ]
    for i, (num, title, sub, is_bible) in enumerate(rows):
        y = 50 + i * 78
        if i == 0:
            d.rectangle([0, y, W, y + 72], fill=(238, 246, 255))
        d.text((14, y + 22), num, fill=(60, 60, 65), font=fR)
        if is_bible:
            d.rounded_rectangle([38, y + 18, 82, y + 38], radius=3,
                                 fill=(187, 222, 251))
            d.text((43, y + 21), "성경", fill=(21, 101, 192), font=fS)
            d.text((90, y + 20), title, fill=(35, 35, 40), font=fR)
        else:
            d.text((38, y + 20), title, fill=(35, 35, 40), font=fR)
        d.text((38, y + 42), sub, fill=(155, 155, 160), font=fS)
        d.text((W - 55, y + 20), "✕   ⠿", fill=(175, 175, 180), font=fR)
        d.line([(0, y + 72), (W, y + 72)], fill=(232, 232, 236))

    d.text((16, 288), "드래그 핸들(⠿)로 순서 변경  |  ✕ 버튼으로 제거",
           fill=(160, 160, 165), font=fS)

    return img


def make_design_panel_image() -> Image.Image:
    """디자인 패널 이미지."""
    W, H = 580, 480
    img = Image.new('RGB', (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)

    fB = _font(14, bold=True)
    fR = _font(13)
    fS = _font(11)

    d.rectangle([0, 0, W, 44], fill=(244, 244, 247))
    d.text((16, 12), "PPTX 디자인", fill=(25, 25, 30), font=fB)
    d.line([(0, 44), (W, 44)], fill=(218, 218, 222))

    # 미리보기 슬라이드
    d.rounded_rectangle([16, 54, W - 16, 230], radius=6, fill=(27, 27, 27))
    d.text((90, 118), "주님 한 분만으로", fill=(255, 255, 255),
           font=_font(20, bold=True))
    d.text((90, 150), "Only You, Lord", fill=(255, 241, 118), font=_font(15))
    d.text((16, 236), "▲ 선택된 첫 번째 곡 · 첫 페이지 실시간 미리보기",
           fill=(125, 125, 130), font=fS)

    # 옵션 행
    rows = [
        ("찬양 글자 크기",  "30 pt  [−] [+]"),
        ("성경 글자 크기",  "30 pt  [−] [+]"),
        ("배경색",         "■  #1B1B1B  ▾"),
        ("찬양 글자색",    "■  #FFFFFF  ▾"),
        ("성경 글자색",    "■  #FFFFFF  ▾"),
        ("영어 가사 포함", "● 켜짐"),
        ("곡 제목 표시",   "꺼짐"),
        ("글자 정렬",      "[좌] [●중] [우]"),
        ("글자 위치",      "[상] [●중] [하]"),
    ]
    for i, (label, val) in enumerate(rows):
        y = 260 + i * 22
        d.text((16, y), label, fill=(90, 90, 96), font=fR)
        d.text((210, y), val, fill=(32, 32, 38), font=fR)

    d.rounded_rectangle([16, 456, W - 16, 486], radius=7, fill=(15, 139, 141))
    d.text((240, 462), "PPTX 저장", fill=(255, 255, 255), font=fB)

    return img


def make_presentation_image() -> Image.Image:
    """발표 모드 컨트롤 이미지."""
    W, H = 920, 230
    img = Image.new('RGB', (W, H), (242, 246, 250))
    d = ImageDraw.Draw(img)

    fB = _font(14, bold=True)
    fR = _font(13)
    fS = _font(11)

    # 컨트롤 바
    d.rounded_rectangle([10, 8, W - 10, 64], radius=8, fill=(196, 228, 232))
    d.text((24, 24), "■  발표 중  |  주님 한 분만으로", fill=(8, 55, 65), font=fB)
    d.text((560, 24), "◀", fill=(8, 55, 65), font=fB)
    d.rounded_rectangle([600, 18, 700, 52], radius=5, fill=(255, 255, 255))
    d.text((626, 26), "3 / 12", fill=(30, 30, 35), font=fB)
    d.text((720, 24), "▶", fill=(8, 55, 65), font=fB)
    d.rounded_rectangle([780, 18, 900, 52], radius=5, fill=(228, 100, 100))
    d.text((808, 26), "■  종료", fill=(255, 255, 255), font=fB)

    # 슬라이드 썸네일 목록
    slide_names = [
        "1. 주님 한 분만으로",
        "2. 주님 한 분만으로",
        "3. 주님 한 분만으로",
        "4. 주의 이름 높이세",
        "5. 로마서 8:28",
    ]
    for i, name in enumerate(slide_names):
        x = 10 + i * 178
        color = (15, 139, 141) if i == 2 else (196, 208, 215)
        tc = (255, 255, 255) if i == 2 else (75, 88, 95)
        d.rounded_rectangle([x, 74, x + 166, 180], radius=6, fill=color)
        d.text((x + 10, 116), name, fill=tc, font=fS)

    d.text((18, 196),
           "키보드:  ← →  이전/다음  |  숫자 + Enter  슬라이드 번호로 이동  |  ESC  발표 종료",
           fill=(100, 118, 128), font=fS)

    return img


def make_bible_image() -> Image.Image:
    """성경 본문 추가 패널 이미지."""
    W, H = 620, 340
    img = Image.new('RGB', (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)

    fB = _font(14, bold=True)
    fR = _font(13)
    fS = _font(11)

    d.rectangle([0, 0, W, 44], fill=(244, 244, 247))
    d.text((16, 12), "성경 본문", fill=(25, 25, 30), font=fB)
    d.line([(0, 44), (W, 44)], fill=(218, 218, 222))

    # 탭
    d.text((16, 58), "찬양 검색", fill=(155, 155, 160), font=fR)
    d.text((140, 58), "성경 본문", fill=(15, 139, 141), font=fB)
    d.line([(140, 76), (240, 76)], fill=(15, 139, 141), width=2)

    # 드롭다운 선택
    labels = ["버전", "책", "장"]
    vals = ["개역개정  ▾", "로마서  ▾", "8장  ▾"]
    for i, (lbl, val) in enumerate(zip(labels, vals)):
        x = 16 + i * 196
        d.text((x, 90), lbl, fill=(105, 105, 110), font=fR)
        d.rounded_rectangle([x, 110, x + 180, 138], radius=4,
                             outline=(185, 185, 190), fill=(251, 251, 253))
        d.text((x + 10, 117), val, fill=(38, 38, 44), font=fR)

    d.line([(16, 152), (W - 16, 152)], fill=(228, 228, 232))

    # 절 목록
    verses = [
        ("28절", "우리가 알거니와 하나님을 사랑하는 자…"),
        ("29절", "하나님이 미리 아신 자들을 또한 그 아들의…"),
        ("30절", "또 미리 정하신 그들을 또한 부르시고…"),
    ]
    for i, (v, text) in enumerate(verses):
        y = 160 + i * 52
        d.text((16, y + 14), v, fill=(12, 98, 115), font=fB)
        d.text((80, y + 14), text, fill=(48, 48, 54), font=fR)
        d.rounded_rectangle([W - 100, y + 8, W - 14, y + 42], radius=5,
                             fill=(15, 139, 141))
        d.text((W - 88, y + 16), "+ 추가", fill=(255, 255, 255), font=fR)
        d.line([(16, y + 50), (W - 16, y + 50)], fill=(238, 238, 242))

    return img


def make_export_flow_image() -> Image.Image:
    """PPTX 내보내기 흐름 이미지."""
    W, H = 900, 160
    img = Image.new('RGB', (W, H), (244, 248, 252))
    d = ImageDraw.Draw(img)

    fB = _font(14, bold=True)
    fR = _font(12)

    steps = [
        ("① 곡/성경 선택", "스테이징에 항목 추가"),
        ("② 디자인 설정", "색상·크기·위치 조정"),
        ("③ PPTX 저장 클릭", "저장 위치 선택"),
        ("④ 완료", "worship_slides.pptx 생성"),
    ]
    cols = [(20, 54, 66), (15, 96, 118), (15, 139, 141), (12, 155, 100)]
    box_w = 200

    for i, ((line1, line2), col) in enumerate(zip(steps, cols)):
        x0 = 16 + i * (box_w + 20)
        x1 = x0 + box_w
        d.rounded_rectangle([x0, 12, x1, 130], radius=9, fill=col)
        d.text((x0 + 12, 40), line1, fill=(255, 255, 255), font=fB)
        d.text((x0 + 12, 68), line2, fill=(185, 225, 225), font=fR)
        if i < 3:
            mx = x1 + 10
            d.polygon([(mx, 66), (mx + 14, 76), (mx, 86)],
                       fill=(98, 148, 155))

    return img


# ── 문서 빌더 ────────────────────────────────────────────────────────────────

def add_section_heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph()
    set_para_spacing(para, before=200, after=80)
    if level == 1:
        run = para.add_run(text)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = DARK_BG
    else:
        run = para.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = ACCENT_TEAL
    return para


def add_body_text(doc: Document, text: str, indent=False):
    para = doc.add_paragraph()
    set_para_spacing(para, before=40, after=40)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT_DARK
    return para


def add_note_box(doc: Document, text: str, box_type='note'):
    """노트/경고/팁 박스."""
    colors = {
        'note': ('E8F4F8', '0F8B8D', '📌 참고'),
        'warn': ('FFF3E0', 'E65100', '⚠ 주의'),
        'tip':  ('E8F5E9', '2E7D32', '💡 팁'),
    }
    bg, border_c, prefix = colors.get(box_type, colors['note'])

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    set_cell_border(cell,
                    left={'val': 'single', 'sz': 24, 'color': border_c},
                    top={'val': 'none', 'sz': 0, 'color': 'FFFFFF'},
                    right={'val': 'none', 'sz': 0, 'color': 'FFFFFF'},
                    bottom={'val': 'none', 'sz': 0, 'color': 'FFFFFF'})
    para = cell.paragraphs[0]
    para.paragraph_format.left_indent = Cm(0.3)
    set_para_spacing(para, before=60, after=60)
    r1 = para.add_run(prefix + '  ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = para.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def add_step_table(doc: Document, steps: list):
    """번호 붙은 단계 표."""
    table = doc.add_table(rows=len(steps), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Cm(1.2), Cm(13.5)]
    for i, row in enumerate(table.rows):
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        num, desc = steps[i]
        # 번호 셀
        set_cell_bg(row.cells[0], '143642')
        para0 = row.cells[0].paragraphs[0]
        para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para0.add_run(str(num))
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = WHITE
        # 내용 셀
        para1 = row.cells[1].paragraphs[0]
        para1.paragraph_format.left_indent = Cm(0.2)
        r2 = para1.add_run(desc)
        r2.font.size = Pt(10.5)
    doc.add_paragraph()


def add_key_table(doc: Document, keys: list):
    """단축키 표."""
    # 헤더
    table = doc.add_table(rows=1 + len(keys), cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for j, title in enumerate(['단축키', '기능']):
        cell = hdr.cells[j]
        set_cell_bg(cell, '0F4C5C')
        para = cell.paragraphs[0]
        r = para.add_run(title)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for i, (key, func) in enumerate(keys):
        row = table.rows[i + 1]
        set_cell_bg(row.cells[0], 'F0F8F8' if i % 2 == 0 else 'FFFFFF')
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(key)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = ACCENT_BLUE
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(func)
        r1.font.size = Pt(10)
    doc.add_paragraph()


def add_image_to_doc(doc: Document, img: Image.Image,
                      caption: str, width_cm=14.5):
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    doc.add_picture(buf, width=Cm(width_cm))
    # 가운데 정렬
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 캡션
    cap = doc.add_paragraph(f"[그림] {caption}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(cap, before=20, after=120)
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    cap.runs[0].font.italic = True


def add_divider(doc: Document):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0F8B8D')
    pBdr.append(bottom)
    pPr.append(pBdr)
    set_para_spacing(para, before=120, after=120)


# ── 커버 페이지 ───────────────────────────────────────────────────────────────

def build_cover(doc: Document):
    # 커버 배경 색상 박스
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover_table.cell(0, 0)
    set_cell_bg(cell, '143642')
    cell.width = Cm(16)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(para, before=600, after=200)

    r_title = para.add_run('\n예배 슬라이드\n사용자 설명서\n')
    r_title.font.size = Pt(28)
    r_title.font.bold = True
    r_title.font.color.rgb = WHITE

    r_sub = para.add_run('\nWorship Slides — 찬양·성경 PPT 관리 도구\n')
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0xA0, 0xD8, 0xD8)

    r_ver = para.add_run('\nv1.0.12   |   2026년 5월 기준\n\n')
    r_ver.font.size = Pt(10)
    r_ver.font.color.rgb = RGBColor(0x70, 0xA0, 0xA0)

    doc.add_paragraph()
    doc.add_page_break()


# ── 목차 ─────────────────────────────────────────────────────────────────────

def build_toc(doc: Document):
    add_section_heading(doc, '목  차', level=1)
    toc_items = [
        ('1', '프로그램 소개', '3'),
        ('2', '처음 시작하기', '3'),
        ('  2.1', '프로그램 실행', '3'),
        ('  2.2', '화면 구성 한눈에 보기', '4'),
        ('3', '찬양 가져오기 (PPT/PPTX 폴더 임포트)', '5'),
        ('4', '찬양 검색 및 선택', '6'),
        ('5', '성경 본문 추가', '7'),
        ('  5.1', '성경 데이터 불러오기 (최초 1회)', '7'),
        ('  5.2', '성경 본문 슬라이드 추가', '7'),
        ('6', '선택 순서 관리 (스테이징)', '8'),
        ('7', '디자인 설정', '9'),
        ('8', 'PPTX 파일 내보내기', '10'),
        ('9', '발표 모드 (화면 발표)', '11'),
        ('10', '곡 직접 추가 / 편집', '13'),
        ('11', '프로그램 업데이트', '13'),
        ('12', '자주 묻는 질문 (FAQ)', '14'),
    ]
    for num, title, page in toc_items:
        para = doc.add_paragraph()
        set_para_spacing(para, before=30, after=30)
        indent = len(num) - len(num.lstrip())
        para.paragraph_format.left_indent = Cm(indent * 0.5)
        r_num = para.add_run(f'{num}  ')
        r_num.font.size = Pt(10.5)
        r_num.font.bold = (not num.startswith(' '))
        r_num.font.color.rgb = ACCENT_TEAL if not num.startswith(' ') else TEXT_DARK
        r_title = para.add_run(title)
        r_title.font.size = Pt(10.5)
        r_title.font.bold = (not num.startswith(' '))

    doc.add_page_break()


# ── 본문 ─────────────────────────────────────────────────────────────────────

def build_body(doc: Document):

    # ── 1. 프로그램 소개 ──────────────────────────────────────────────────
    add_section_heading(doc, '1. 프로그램 소개', 1)
    add_body_text(doc,
        '예배 슬라이드(Worship Slides)는 교회 예배 준비에 특화된 데스크탑 전용 도구입니다. '
        'PPT/PPTX 형식으로 보관된 찬양 파일을 한 곳에 모아 관리하고, '
        '원하는 곡과 성경 본문만 골라 새 PPTX로 내보내거나 '
        '외부 모니터(스크린)에 바로 발표할 수 있습니다.')

    # 주요 기능 표
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    headers = [('주요 기능', '설명')]
    rows_data = [
        ('폴더 일괄 가져오기', '.ppt / .pptx 파일을 폴더째 불러와 자동으로 데이터베이스에 저장'),
        ('한/영 가사 자동 분리', '슬라이드 내 한국어와 영어가 섞여 있어도 자동으로 구분하여 저장'),
        ('실시간 검색', '곡명 또는 가사 내용으로 즉시 검색'),
        ('성경 본문 슬라이드', 'JSON 형식의 성경 데이터를 불러와 원하는 절을 슬라이드로 추가'),
        ('PPTX 내보내기', '배경색·글자색·크기·위치 등 스타일을 직접 설정해 PPTX 생성'),
        ('화면 발표 모드', '외부 스크린에 슬라이드를 전체화면으로 송출, 키보드로 페이지 이동'),
    ]
    hdr_row = table.rows[0]
    for j, (h, _) in enumerate([headers[0]]):
        pass
    for j, title in enumerate(['주요 기능', '설명']):
        set_cell_bg(table.rows[0].cells[j], '0F4C5C')
        r = table.rows[0].cells[j].paragraphs[0].add_run(title)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for i, (feat, desc) in enumerate(rows_data):
        row = table.rows[i + 1]
        set_cell_bg(row.cells[0], 'E8F4F8' if i % 2 == 0 else 'FFFFFF')
        r0 = row.cells[0].paragraphs[0].add_run(feat)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = ACCENT_BLUE
        r1 = row.cells[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(10)
    doc.add_paragraph()
    add_divider(doc)

    # ── 2. 처음 시작하기 ──────────────────────────────────────────────────
    add_section_heading(doc, '2. 처음 시작하기', 1)

    add_section_heading(doc, '2.1  프로그램 실행', 2)
    add_body_text(doc,
        '설치 폴더에서 worship_slides.exe (Windows) 또는 '
        'Worship Slides.app (macOS)를 실행합니다. '
        '첫 실행 시 데이터베이스가 자동으로 초기화되며 '
        '바로 사용할 수 있는 상태로 시작됩니다.')
    add_note_box(doc,
        '.ppt(구형) 파일을 가져오려면 LibreOffice가 설치되어 있어야 합니다. '
        '.pptx 파일은 LibreOffice 없이도 정상적으로 불러올 수 있습니다.',
        'note')

    add_section_heading(doc, '2.2  화면 구성 한눈에 보기', 2)
    add_body_text(doc, '프로그램을 실행하면 다음과 같은 화면이 표시됩니다.')
    add_image_to_doc(doc, make_main_ui_image(),
                      '메인 화면 구성 — 상단 바, 왼쪽 작업 영역, 오른쪽 디자인 패널')

    # 구성 요소 설명 표
    parts = [
        ('상단 헤더 바', '찬양폴더 선택, 성경 불러오기, 저장된 곡 수 표시, 업데이트 확인'),
        ('발표 컨트롤 바', '발표 시작/종료, 슬라이드 이전/다음, 현재 위치 표시'),
        ('선택한 순서 패널', '발표할 항목(찬양·성경)의 순서 목록. 드래그로 순서 변경 가능'),
        ('찬양 검색 패널', '검색창과 곡 목록. 체크하면 "선택한 순서"에 자동 추가'),
        ('성경 본문 탭', '책·장·절을 선택해 성경 슬라이드를 순서에 추가'),
        ('PPTX 디자인 패널', '슬라이드 미리보기와 스타일 설정, PPTX 저장 버튼'),
    ]
    table2 = doc.add_table(rows=len(parts) + 1, cols=2)
    table2.style = 'Table Grid'
    for j, title in enumerate(['구성 요소', '역할']):
        set_cell_bg(table2.rows[0].cells[j], '143642')
        r = table2.rows[0].cells[j].paragraphs[0].add_run(title)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for i, (part, role) in enumerate(parts):
        row = table2.rows[i + 1]
        set_cell_bg(row.cells[0], 'F0F8F8' if i % 2 == 0 else 'FFFFFF')
        r0 = row.cells[0].paragraphs[0].add_run(part)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = ACCENT_BLUE
        r1 = row.cells[1].paragraphs[0].add_run(role)
        r1.font.size = Pt(10)
    doc.add_paragraph()
    add_divider(doc)

    # ── 3. 찬양 가져오기 ─────────────────────────────────────────────────
    add_section_heading(doc, '3. 찬양 가져오기 (PPT/PPTX 폴더 임포트)', 1)
    add_body_text(doc,
        '예배에서 사용할 찬양 PPT/PPTX 파일이 저장된 폴더를 프로그램에 불러옵니다. '
        '한 번 가져온 곡은 데이터베이스에 저장되므로, '
        '이후에는 앱을 재시작해도 다시 가져올 필요가 없습니다.')

    add_image_to_doc(doc, make_import_image(),
                      '찬양 가져오기 처리 흐름')

    add_step_table(doc, [
        (1, '상단 헤더 바의 [찬양폴더 선택] 버튼을 클릭합니다.'),
        (2, '폴더 선택 대화상자에서 PPT/PPTX 파일이 있는 폴더를 선택합니다.\n'
            '      하위 폴더까지 자동으로 탐색합니다.'),
        (3, '읽는 중… 상태가 표시되며 파일을 분석합니다. 파일 수에 따라 수 초~수 분 소요됩니다.'),
        (4, '완료 후 추가된 곡 수를 알림으로 확인합니다.\n'
            '      이미 저장된 곡은 자동으로 건너뛰어 중복이 방지됩니다.'),
    ])

    add_note_box(doc,
        '같은 폴더를 여러 번 가져와도 안전합니다. '
        '이미 저장된 곡은 "중복 N개 건너뜀"으로 표시되고 새 곡만 추가됩니다.',
        'tip')
    add_note_box(doc,
        '.ppt(구형) 파일이 포함된 폴더를 선택했을 때 LibreOffice가 없으면 '
        '"LibreOffice가 필요합니다" 안내창이 표시됩니다. '
        '.pptx 파일만 있다면 이 안내는 나타나지 않습니다.',
        'warn')
    add_divider(doc)

    # ── 4. 찬양 검색 및 선택 ────────────────────────────────────────────
    add_section_heading(doc, '4. 찬양 검색 및 선택', 1)
    add_body_text(doc,
        '검색창에 곡명이나 가사 일부를 입력하면 저장된 곡 중에서 '
        '일치하는 항목이 실시간으로 표시됩니다.')

    add_step_table(doc, [
        (1, '왼쪽 하단 [찬양 검색] 패널의 검색창에 곡명 또는 가사 키워드를 입력합니다.'),
        (2, '목록에서 원하는 곡을 클릭해 체크(✓) 표시합니다.\n'
            '      체크된 곡은 즉시 "선택한 순서" 패널에 추가됩니다.'),
        (3, '동일한 곡을 다시 클릭하면 체크가 해제되고 순서 목록에서도 제거됩니다.'),
        (4, '검색창을 비우면 전체 곡 목록이 표시됩니다.'),
    ])

    add_note_box(doc,
        '곡 목록 하단의 […] 메뉴를 통해 선택한 곡만 삭제하거나 '
        'DB 전체를 초기화할 수 있습니다.',
        'note')
    add_divider(doc)

    # ── 5. 성경 본문 추가 ────────────────────────────────────────────────
    add_section_heading(doc, '5. 성경 본문 추가', 1)

    add_section_heading(doc, '5.1  성경 데이터 불러오기 (최초 1회)', 2)
    add_body_text(doc,
        '성경 본문 탭을 사용하려면 먼저 성경 JSON 파일을 불러와야 합니다. '
        '이 작업은 처음 한 번만 하면 이후에는 저장된 데이터가 유지됩니다.')
    add_step_table(doc, [
        (1, '상단 헤더 바의 [성경 불러오기] 버튼을 클릭합니다.'),
        (2, '성경 JSON 파일을 선택합니다. (개역개정, 새번역 등 다양한 형식 지원)'),
        (3, '"버전 이름" 입력창에 예: 개역개정, 새번역, KJV 등을 입력하고 [저장]을 클릭합니다.'),
        (4, '"개역개정 성경 N절을 저장했습니다" 알림이 뜨면 완료입니다.'),
    ])
    add_note_box(doc,
        '지원하는 JSON 형식은 세 가지입니다:\n'
        '• 형식 A: [{"book":"창세기","chapter":1,"verse":1,"text":"…"}]\n'
        '• 형식 B: [{"book":"창세기","chapters":[…]}]\n'
        '• 형식 C: {"창세기":{"1":{"1":"태초에…"}}}',
        'note')

    add_section_heading(doc, '5.2  성경 본문 슬라이드 추가', 2)
    add_image_to_doc(doc, make_bible_image(),
                      '성경 본문 탭 — 버전·책·장 선택 후 절 추가')
    add_step_table(doc, [
        (1, '[찬양 검색] 패널 상단의 [성경 본문] 탭을 클릭합니다.'),
        (2, '버전 · 책 · 장을 드롭다운으로 선택합니다.'),
        (3, '표시된 절 목록에서 원하는 절 오른쪽의 [+ 추가] 버튼을 클릭합니다.\n'
            '      해당 절이 "선택한 순서" 패널에 성경 항목으로 추가됩니다.'),
    ])
    add_divider(doc)

    # ── 6. 선택 순서 관리 ────────────────────────────────────────────────
    add_section_heading(doc, '6. 선택 순서 관리 (스테이징)', 1)
    add_body_text(doc,
        '"선택한 순서" 패널은 실제 발표하거나 내보낼 찬양·성경의 목록과 순서를 관리합니다.')
    add_image_to_doc(doc, make_staging_image(),
                      '"선택한 순서" 패널 — 찬양 2곡과 성경 1절이 추가된 모습')

    staging_funcs = [
        ('순서 변경', '각 항목 오른쪽의 ⠿ 핸들을 드래그하여 위아래로 이동합니다.'),
        ('항목 제거', '항목 오른쪽의 ✕ 버튼을 클릭하면 순서 목록에서 제거됩니다.\n'
                      '       찬양의 경우 검색 목록의 체크도 자동 해제됩니다.'),
        ('미리보기 전환', '항목을 클릭하면 오른쪽 디자인 패널의 미리보기가 해당 곡으로 전환됩니다.'),
        ('패널 접기', '패널 상단의 [접기] 아이콘으로 패널을 최소화해 화면 공간을 넓힐 수 있습니다.'),
        ('패널 크기 조절', '두 패널 사이의 구분선을 드래그하면 상하 비율을 조절할 수 있습니다.'),
    ]
    table3 = doc.add_table(rows=len(staging_funcs) + 1, cols=2)
    table3.style = 'Table Grid'
    for j, title in enumerate(['기능', '방법']):
        set_cell_bg(table3.rows[0].cells[j], '0F4C5C')
        r = table3.rows[0].cells[j].paragraphs[0].add_run(title)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for i, (func, how) in enumerate(staging_funcs):
        row = table3.rows[i + 1]
        set_cell_bg(row.cells[0], 'E8F4F8' if i % 2 == 0 else 'FFFFFF')
        r0 = row.cells[0].paragraphs[0].add_run(func)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = ACCENT_BLUE
        r1 = row.cells[1].paragraphs[0].add_run(how)
        r1.font.size = Pt(10)
    doc.add_paragraph()
    add_divider(doc)

    # ── 7. 디자인 설정 ───────────────────────────────────────────────────
    add_section_heading(doc, '7. 디자인 설정', 1)
    add_body_text(doc,
        '오른쪽 "PPTX 디자인" 패널에서 슬라이드의 시각적 스타일을 설정합니다. '
        '변경 사항은 상단 미리보기에 즉시 반영됩니다.')
    add_image_to_doc(doc, make_design_panel_image(),
                      'PPTX 디자인 패널 — 미리보기와 스타일 옵션')

    design_opts = [
        ('찬양 글자 크기', '찬양 가사의 글자 크기(pt). 기본값: 30'),
        ('성경 글자 크기', '성경 본문의 글자 크기(pt). 기본값: 30'),
        ('배경색', '슬라이드 배경색. 스와치(팔레트) 또는 직접 입력 가능'),
        ('찬양 글자색', '찬양 가사 글자색'),
        ('성경 글자색', '성경 본문 글자색'),
        ('영어 가사 포함', '켜면 영어 가사가 함께 표시됨. 끄면 한국어만 표시'),
        ('영어 글자색', '영어 가사의 글자색. 기본값: 노랑(#FFF176)'),
        ('곡 제목 표시', '켜면 슬라이드 모서리에 곡 이름을 작은 글씨로 표시'),
        ('성경 제목 표시', '켜면 슬라이드 모서리에 성경 참조(예: 롬 8:28)를 표시'),
        ('글자 정렬', '가사 수평 정렬: 좌측 / 중앙 / 우측'),
        ('글자 위치', '슬라이드 내 가사 수직 위치: 상단 / 중단 / 하단'),
    ]
    table4 = doc.add_table(rows=len(design_opts) + 1, cols=2)
    table4.style = 'Table Grid'
    for j, title in enumerate(['설정 항목', '설명']):
        set_cell_bg(table4.rows[0].cells[j], '143642')
        r = table4.rows[0].cells[j].paragraphs[0].add_run(title)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
    for i, (opt, desc) in enumerate(design_opts):
        row = table4.rows[i + 1]
        set_cell_bg(row.cells[0], 'E8F4F8' if i % 2 == 0 else 'FFFFFF')
        r0 = row.cells[0].paragraphs[0].add_run(opt)
        r0.font.bold = True
        r0.font.size = Pt(10)
        r0.font.color.rgb = ACCENT_BLUE
        r1 = row.cells[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(10)
    doc.add_paragraph()
    add_note_box(doc,
        '설정된 스타일은 앱 재시작 후에도 자동으로 복원됩니다. '
        '별도 저장 없이도 마지막 설정이 유지됩니다.',
        'tip')
    add_divider(doc)

    # ── 8. PPTX 내보내기 ────────────────────────────────────────────────
    add_section_heading(doc, '8. PPTX 파일 내보내기', 1)
    add_body_text(doc,
        '"선택한 순서"에 담긴 찬양과 성경 본문을 하나의 PPTX 파일로 내보냅니다. '
        '생성된 파일은 PowerPoint, Keynote, LibreOffice Impress 등에서 바로 열 수 있습니다.')
    add_image_to_doc(doc, make_export_flow_image(),
                      'PPTX 내보내기 흐름')
    add_step_table(doc, [
        (1, '"선택한 순서" 패널에 내보낼 찬양·성경이 모두 추가되어 있는지 확인합니다.'),
        (2, '"PPTX 디자인" 패널에서 미리보기를 확인하며 스타일을 조정합니다.'),
        (3, '[PPTX 저장] 버튼을 클릭합니다.'),
        (4, '저장 대화상자에서 파일 이름과 저장 위치를 선택한 뒤 [저장]을 클릭합니다.'),
        (5, '"PPTX 저장 완료: /경로/파일명.pptx" 알림이 표시되면 완료입니다.'),
    ])
    add_note_box(doc,
        '파일 이름을 입력할 때 확장자(.pptx)를 생략해도 자동으로 추가됩니다.',
        'tip')
    add_divider(doc)

    # ── 9. 발표 모드 ─────────────────────────────────────────────────────
    add_section_heading(doc, '9. 발표 모드 (화면 발표)', 1)
    add_body_text(doc,
        '발표 모드를 시작하면 외부 모니터(빔 프로젝터, 스크린 등)에 '
        '슬라이드가 전체화면으로 표시됩니다. '
        '모니터가 하나뿐인 경우에도 별도 창으로 열립니다.')
    add_image_to_doc(doc, make_presentation_image(),
                      '발표 컨트롤 바 — 발표 중 화면')

    add_section_heading(doc, '발표 시작 방법', 2)
    add_step_table(doc, [
        (1, '"선택한 순서"에 항목이 하나 이상 있는 상태에서\n'
            '      발표 컨트롤 바의 [▶ 발표 시작] 버튼을 클릭합니다.'),
        (2, '외부 모니터 또는 새 창에 첫 번째 슬라이드가 표시됩니다.'),
        (3, '키보드 또는 화면의 버튼으로 슬라이드를 넘깁니다.'),
        (4, '[■ 종료] 버튼 또는 ESC 키를 누르면 발표를 종료합니다.'),
    ])

    add_section_heading(doc, '발표 중 슬라이드 제어', 2)
    add_key_table(doc, [
        ('→  또는  ↓', '다음 슬라이드로 이동'),
        ('←  또는  ↑', '이전 슬라이드로 이동'),
        ('숫자 + Enter', '해당 번호 슬라이드로 바로 이동 (예: 5 → Enter → 5번 슬라이드)'),
        ('ESC', '슬라이드 번호 입력 취소 / 발표 종료'),
    ])

    add_section_heading(doc, '발표 중 슬라이드 편집', 2)
    add_body_text(doc,
        '발표 중에도 슬라이드 내용을 즉석에서 수정할 수 있습니다.')
    add_step_table(doc, [
        (1, '발표 컨트롤 바 아래의 슬라이드 목록 패널에서 수정할 슬라이드를 더블클릭합니다.'),
        (2, '편집 대화상자에서 한국어 가사 및 영어 가사를 수정합니다.'),
        (3, '[저장]을 클릭하면 발표 화면에 즉시 반영됩니다.'),
    ])
    add_note_box(doc,
        '편집 내용은 현재 세션에만 적용됩니다. '
        '데이터베이스에 영구 저장하려면 곡 편집 기능을 사용하세요.',
        'note')
    add_divider(doc)

    # ── 10. 곡 직접 추가 / 편집 ─────────────────────────────────────────
    add_section_heading(doc, '10. 곡 직접 추가 / 편집', 1)
    add_body_text(doc,
        'PPT 파일 없이 새 곡을 직접 입력하거나 저장된 곡의 가사를 수정할 수 있습니다.')

    add_step_table(doc, [
        (1, '검색 패널 상단의 [+ 곡 추가] 버튼을 클릭해 새 곡을 입력합니다.\n'
            '      기존 곡을 수정하려면 목록에서 곡을 마우스 오른쪽 클릭 → [편집]을 선택합니다.'),
        (2, '제목과 가사를 입력합니다. 페이지 구분은 "###"을 줄 단독으로 입력합니다.\n'
            '      예) 주님 한 분만으로\n'
            '           ###\n'
            '           내게는 족해'),
        (3, '[저장]을 클릭하면 데이터베이스에 영구 저장됩니다.'),
    ])
    add_divider(doc)

    # ── 11. 업데이트 ─────────────────────────────────────────────────────
    add_section_heading(doc, '11. 프로그램 업데이트', 1)
    add_body_text(doc,
        '프로그램 시작 시 자동으로 최신 버전을 확인합니다. '
        '새 버전이 있으면 상단에 업데이트 안내 배너가 표시됩니다.')
    add_step_table(doc, [
        (1, '업데이트 배너의 [업데이트] 버튼을 클릭합니다.'),
        (2, '다운로드 진행 표시줄이 표시됩니다. 완료되면 앱이 자동으로 재시작됩니다.'),
    ])
    add_note_box(doc,
        '수동으로 업데이트를 확인하려면 상단 헤더 바의 [↺] 버튼을 클릭합니다.',
        'tip')
    add_divider(doc)

    # ── 12. FAQ ──────────────────────────────────────────────────────────
    add_section_heading(doc, '12. 자주 묻는 질문 (FAQ)', 1)

    faqs = [
        ('Q. 찬양 가져오기 후 곡이 하나도 추가되지 않습니다.',
         'A. 선택한 폴더에 .pptx 또는 .ppt 파일이 없을 수 있습니다. '
         '올바른 폴더를 선택했는지 확인해 주세요. '
         '.ppt 파일만 있는 경우 LibreOffice 설치 여부도 확인하세요.'),
        ('Q. "중복 N개 건너뜀" 메시지가 표시됩니다.',
         'A. 이미 DB에 저장된 곡이 포함된 폴더를 다시 선택한 경우입니다. '
         '정상 동작으로, 새 곡만 추가됩니다.'),
        ('Q. 발표 화면이 주 모니터에 나타납니다.',
         'A. 운영체제의 디스플레이 설정에서 외부 모니터가 "확장" 모드인지 확인하세요. '
         '복제 모드에서는 보조 모니터에 전체화면 표시가 불가할 수 있습니다.'),
        ('Q. 성경 불러오기 버튼을 눌렀는데 절이 표시되지 않습니다.',
         'A. JSON 파일 형식이 지원하는 세 가지 형식 중 하나인지 확인하세요. '
         '5.1절의 형식 설명을 참고하세요.'),
        ('Q. 영어 가사가 없는데 빈 영어 줄이 표시됩니다.',
         'A. "PPTX 디자인" 패널에서 [영어 가사 포함]을 꺼주세요.'),
        ('Q. PPTX 저장 후 PowerPoint에서 글꼴이 다르게 보입니다.',
         'A. 내보내기 시 기본 글꼴(Pretendard)이 해당 PC에 설치되지 않으면 '
         '대체 글꼴로 표시될 수 있습니다.'),
        ('Q. 저장된 곡을 모두 삭제하고 다시 시작하고 싶습니다.',
         'A. 검색 패널 메뉴의 [전체 초기화] 버튼을 클릭하세요. '
         '삭제 전 확인 대화상자가 표시됩니다.'),
    ]

    for q, a in faqs:
        # 질문
        q_para = doc.add_paragraph()
        set_para_spacing(q_para, before=80, after=20)
        q_para.paragraph_format.left_indent = Cm(0)
        r_q = q_para.add_run(q)
        r_q.font.bold = True
        r_q.font.size = Pt(10.5)
        r_q.font.color.rgb = ACCENT_BLUE
        # 답변
        a_para = doc.add_paragraph()
        set_para_spacing(a_para, before=0, after=60)
        a_para.paragraph_format.left_indent = Cm(0.5)
        r_a = a_para.add_run(a)
        r_a.font.size = Pt(10.5)
        r_a.font.color.rgb = TEXT_DARK

    add_divider(doc)

    # 꼬리말
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(footer_para, before=200)
    r_f = footer_para.add_run('예배 슬라이드 (Worship Slides)  |  문의 및 최신 버전: github.com/yjchae/make_ppt')
    r_f.font.size = Pt(9)
    r_f.font.color.rgb = RGBColor(0xA0, 0xA0, 0xA0)
    r_f.font.italic = True


# ── 문서 페이지 설정 ─────────────────────────────────────────────────────────

def setup_page(doc: Document):
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)


# ── 메인 ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    setup_page(doc)

    # 기본 본문 스타일 조정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)

    build_cover(doc)
    build_toc(doc)
    build_body(doc)

    out_path = '/home/user/make_ppt/예배슬라이드_사용자설명서.docx'
    doc.save(out_path)
    print(f'저장 완료: {out_path}')


if __name__ == '__main__':
    main()
